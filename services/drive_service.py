import logging
from io import BytesIO
from core.logger import logger
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload, MediaIoBaseUpload
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from google.oauth2 import service_account
from core.utils import get_env_file_path, safe_execute
import os
from dotenv import load_dotenv
import tempfile
from docx import Document
from docx.shared import Pt
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
import pickle
from typing import List, Dict
import base64
import requests

load_dotenv()

SERVICE_ACCOUNT_JSON = get_env_file_path("SERVICE_ACCOUNT_FILE")
# Права доступа
SCOPES = ['https://www.googleapis.com/auth/drive.file']

# Путь к OAuth JSON
CREDENTIALS_FILE = get_env_file_path("OAUTH_ACCOUNT_FILE")
CORE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "core"))
TOKEN_FILE = os.path.join(CORE_DIR, "token.pickle")
logger.info(TOKEN_FILE)


def get_drive_service():
    """Создаём клиент Google Drive API"""
    try:
        creds = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_JSON,
        scopes=["https://www.googleapis.com/auth/drive"]
        )
        service = build("drive", "v3", credentials=creds)
        return service
    except Exception as e:
        logger.error(f"[Drive] Ошибка создания сервиса: {e}")
        return None


def list_files_in_folder(service, folder_id: str):
    """Получить список файлов в папке Google Drive, используя существующий сервис"""
    try:
        if not service:
            logging.error("[Drive] Сервис не инициализирован")
            return []
        query = f"'{folder_id}' in parents and trashed = false"
        results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()
        return results.get("files", [])
    except Exception as e:
        logger.error(f"[Drive] Ошибка при получении файлов: {e}")
        return []


def find_matching_transcription(service, transcription_folder_id: str, base_filename: str):
    """
    Ищет файл в папке транскрипций, у которого совпадает имя без расширения.
    """
    files = safe_execute(list_files_in_folder, service, transcription_folder_id)
    if not files:
        return None

    for f in files:
        mime = f.get("mimeType", "")
        name = f.get("name", "").lower()
        name_without_ext = os.path.splitext(f["name"])[0]
        if mime != "text/vtt" and not name.endswith(".vtt"):
            continue
        if name_without_ext == base_filename:
            return f
    return None

def get_file_link(file_id: str) -> str:
    """
    Возвращает прямую ссылку на файл Google Drive по его ID
    """
    return f"https://drive.google.com/file/d/{file_id}/view?usp=sharing"



def get_drive_service_oauth2():
    """Создаёт клиент Google Drive через OAuth 2.0"""

    creds = None

    # Load existing token
    if os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE, 'rb') as token:
            creds = pickle.load(token)

    # Refresh if expired
    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
        except Exception:
            creds = None

    # Request new token if needed
    if not creds or not creds.valid:
        flow = InstalledAppFlow.from_client_secrets_file(
            CREDENTIALS_FILE,
            SCOPES
        )
        creds = flow.run_local_server(
            port=0,
            access_type="offline",
            prompt="consent"
        )

        with open(TOKEN_FILE, 'wb') as token:
            pickle.dump(creds, token)

    service = build('drive', 'v3', credentials=creds)
    return service


def download_file_to_path(file_id: str, destination_path: str):
    """Скачивает файл с Google Drive по ID в указанный путь"""
    try:
        # Создаём директорию, если её нет
        os.makedirs(os.path.dirname(destination_path), exist_ok=True)
        service = get_drive_service()
        request = service.files().get_media(fileId=file_id)
        with open(destination_path, "wb") as f:
            downloader = MediaIoBaseDownload(f, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
        return True
    except Exception as e:
        logger.error(f"[Drive] Ошибка скачивания файла {file_id}: {e}")
        return False





# def save_transcription_to_drive(speaker_text: List[Dict], folder_id: str, base_filename: str = None):
#     """
#     Сохраняет транскрипцию со спикерами в временный DOCX и загружает в Google Drive.
#
#     Параметры:
#         speaker_text: List[Dict] с ключами 'start', 'end', 'speaker', 'text'
#         folder_id: ID папки Google Drive
#         base_filename: имя файла, если None — auto генерируется
#
#     Возвращает:
#         dict с 'file_id' и 'webViewLink'
#     """
#     try:
#         file_name = f"transcription_{base_filename}.docx"
#
#         # Создаём временный DOCX
#         doc = Document()
#         doc.add_heading("Transcription", level=1)
#
#         for segment in speaker_text:
#             start = round(segment['start'], 2)
#             end = round(segment['end'], 2)
#             speaker = segment['speaker']
#             text = segment['text']
#             p = doc.add_paragraph()
#             run = p.add_run(f"[{start}-{end}] {speaker}: {text}")
#             run.font.size = Pt(11)
#
#         file_stream = BytesIO()
#         doc.save(file_stream)
#         doc = None
#         # Перемещаем курсор в начало
#         file_stream.seek(0)
#         logger.info(f"[Transcription] Временный DOCX создан: {file_name}")
#
#         # Загружаем в Google Drive
#         service = get_drive_service_oauth2()
#
#         file_metadata = {"name": file_name, "parents": [folder_id]}
#         media = MediaIoBaseUpload(file_stream,mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#         file = service.files().create(body=file_metadata, media_body=media, fields="id, webViewLink").execute()
#         logger.info(f"[Drive] Файл '{file_name}' загружен: {file.get('webViewLink')}")
#
#         file_stream.close()  # закрываем поток
#
#
#         return {"file_id": file["id"], "webViewLink": file.get("webViewLink")}
#
#     except Exception as e:
#         logger.error(f"[Drive] Ошибка при сохранении транскрипции: {e}")
#         return None

# Настройки для Apps Script
APPS_SCRIPT_URL = os.getenv("APPS_SCRIPT_URL")
SECRET_KEY = os.getenv("SECRET_KEY")
def save_transcription_to_drive(speaker_text, folder_id, base_filename=None):
    """
    Сохраняет транскрипцию со спикерами в DOCX и загружает в Google Drive через Apps Script.
    Параметры:
        speaker_text: список секций {'start', 'end', 'speaker', 'text'}
        folder_id: ID папки для записи (обычно ключ папки, который использует Apps Script)
        base_filename: имя файла без расширения
    """
    try:
        file_name = f"transcription_{base_filename or 'auto'}.docx"

        # --- Генерация DOCX в памяти ---
        doc = Document()
        doc.add_heading("Transcription", level=1)

        for segment in speaker_text:
            start = round(segment['start'], 2)
            end = round(segment['end'], 2)
            speaker = segment['speaker']
            text = segment['text']

            p = doc.add_paragraph()
            run = p.add_run(f"[{start}-{end}] {speaker}: {text}")
            run.font.size = Pt(11)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        doc = None

        logger.info(f"[Transcription] Временный DOCX создан: {file_name}")

        # --- Кодируем в Base64 ---
        file_b64 = base64.b64encode(file_stream.read()).decode("utf-8")
        file_stream.close()

        # --- Отправляем файл в Apps Script ---
        payload = {
            "secret": SECRET_KEY,
            "folder": folder_id,        # ключ папки в Apps Script, не raw Drive ID
            "name": file_name,
            "content_b64": file_b64
        }

        response = requests.post(APPS_SCRIPT_URL, json=payload)
        result = response.json()

        if not result.get("success"):
            raise Exception(result.get("error", "Unknown error"))

        logger.info(f"[Drive] Файл '{file_name}' загружен: {result.get('url')}")

        return {
            "file_id": result.get("fileId"),
            "webViewLink": result.get("url"),
        }

    except Exception as e:
        logger.error(f"[Drive] Ошибка при сохранении транскрипции: {e}")
        return None